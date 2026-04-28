"""
Acta de Entrega de Servicio — DataCom S.A.  RE-SIGC-SI-AS-1.0

Flujo:
  1. python-docx abre el template original y rellena todos los campos directamente
  2. Para tablas variables (servicios, equipos) copia filas según los datos
  3. LibreOffice convierte el .docx resultante a PDF (headless)
"""
import os
import glob
import subprocess
import tempfile
from copy import deepcopy
from datetime import datetime, timezone

_MESES = {
    "January": "enero",   "February": "febrero", "March": "marzo",
    "April": "abril",     "May": "mayo",          "June": "junio",
    "July": "julio",      "August": "agosto",     "September": "septiembre",
    "October": "octubre", "November": "noviembre","December": "diciembre",
}

_ASSETS   = os.path.normpath(
    os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'assets')
)
_TEMPLATE = os.path.join(_ASSETS, 'template_acta_servicio.docx')   # original limpio


def _fecha_larga(dt: datetime) -> str:
    mes_en = dt.strftime("%B")
    return f"{dt.day} de {_MESES.get(mes_en, mes_en)} de {dt.year}"


# ── Helpers python-docx ───────────────────────────────────────────────────────

def _replace_in_para(paragraph, old: str, new: str) -> bool:
    """Reemplaza texto en un párrafo consolidando sus runs."""
    full = "".join(r.text for r in paragraph.runs)
    if old not in full:
        return False
    for i, run in enumerate(paragraph.runs):
        run.text = full.replace(old, new) if i == 0 else ""
    return True


def _set_cell(cell, text: str, align=None, compact=False):
    """Escribe texto en la celda; líneas separadas por \\n → párrafos.

    align:   valor de WD_ALIGN_PARAGRAPH; None = preserva alineación original.
    compact: si True, elimina espacio antes/después de párrafo y usa interlineado 1.15.
    """
    from docx.shared import Pt
    lines = str(text).split("\n")

    # Preservar alineación original del primer párrafo del template
    orig_align = cell.paragraphs[0].paragraph_format.alignment if cell.paragraphs else None
    effective_align = align if align is not None else orig_align

    def _fmt(p):
        p.paragraph_format.alignment = effective_align
        if compact:
            p.paragraph_format.space_before  = Pt(0)
            p.paragraph_format.space_after   = Pt(0)
            p.paragraph_format.line_spacing  = 1.15

    # Limpiar runs de todos los párrafos existentes
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""

    # Primera línea → primer párrafo (conserva formato original)
    p0 = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if p0.runs:
        p0.runs[0].text = lines[0]
    else:
        p0.add_run(lines[0])
    _fmt(p0)

    # Líneas adicionales → nuevos párrafos con misma alineación y espaciado
    for line in lines[1:]:
        _fmt(cell.add_paragraph(line))


def _fill_table_rows(table, items: list, fields: list):
    """
    Rellena la fila de datos (índice 1) con items[0] y añade filas extra
    copiando la estructura XML de la fila plantilla.
    """
    if not items:
        return

    data_row = table.rows[1]

    # Rellenar primera fila
    for col_idx, field in enumerate(fields):
        if col_idx < len(data_row.cells):
            _set_cell(data_row.cells[col_idx], str(items[0].get(field, "—")))

    # Filas adicionales: clonar XML de la fila plantilla
    for item in items[1:]:
        new_tr = deepcopy(data_row._tr)
        table._tbl.append(new_tr)
        new_row = table.rows[-1]
        for col_idx, field in enumerate(fields):
            if col_idx < len(new_row.cells):
                _set_cell(new_row.cells[col_idx], str(item.get(field, "—")))


def _keep_section_together(doc, heading_text: str, table):
    """
    Mueve toda la sección de firmas (encabezado + texto de conformidad + tabla)
    a la página siguiente como un bloque si no cabe en la página actual.

    Estrategia:
    - Recorre los elementos XML directos del body en orden para encontrar
      todos los párrafos entre el encabezado y la tabla y les aplica
      keep_with_next, formando una cadena ininterrumpida.
    - cantSplit en cada fila evita que una fila se parta internamente.
    - keep_with_next en el primer párrafo de cada celda (excepto última fila)
      encadena las filas entre sí.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _apply_kwn(para_el):
        """Añade w:keepNext al w:pPr del párrafo XML dado."""
        pPr = para_el.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            para_el.insert(0, pPr)
        kwn = OxmlElement('w:keepNext')
        kwn.set(qn('w:val'), '1')
        pPr.append(kwn)

    # Recorrer body en orden XML: encadenar todos los párrafos desde el
    # encabezado hasta (sin incluir) la tabla de firmas.
    heading_found = False
    for el in doc.element.body:
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag == 'p':
            text = ''.join(t.text for t in el.iter(qn('w:t')) if t.text)
            if heading_text in text:
                heading_found = True
            if heading_found:
                _apply_kwn(el)
        elif tag == 'tbl':
            if heading_found:
                break  # llegamos a la tabla de firmas; detenemos la cadena

    # cantSplit + keep_with_next en filas de la tabla
    rows = list(table.rows)
    for r_idx, row in enumerate(rows):
        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), '1')
        trPr.append(cantSplit)
        if r_idx < len(rows) - 1:
            for cell in row.cells:
                if cell.paragraphs:
                    cell.paragraphs[0].paragraph_format.keep_with_next = True


def _replace_in_section(doc, old: str, new: str):
    """Reemplaza texto en cuerpo, encabezados y pies de página."""
    for p in doc.paragraphs:
        _replace_in_para(p, old, new)
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            if hdr:
                for p in hdr.paragraphs:
                    _replace_in_para(p, old, new)
                for tbl in hdr.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                _replace_in_para(p, old, new)
        for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
            if ftr:
                for p in ftr.paragraphs:
                    _replace_in_para(p, old, new)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_para(p, old, new)


# ── Función principal ─────────────────────────────────────────────────────────

def generate_facility_pdf(facility) -> bytes:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Instale python-docx: pip install python-docx")

    # ── Fechas ───────────────────────────────────────────────────────────────
    fecha_doc       = facility.fecha_fin if facility.fecha_fin else datetime.now(timezone.utc)
    fecha_fin_larga = _fecha_larga(fecha_doc)
    fecha_fin_corta = fecha_doc.strftime("%d/%m/%Y")

    # ── Datos ────────────────────────────────────────────────────────────────
    cliente_nombre = "—"
    if facility.cliente:
        try:
            cliente_nombre = facility.cliente.nombre_cliente
        except Exception:
            pass

    tecnico_nombre = "—"
    if facility.tecnico:
        try:
            tecnico_nombre = facility.tecnico.username
        except Exception:
            pass

    direccion = getattr(facility, "direccion_instalacion", "") or "—"
    ot_code   = facility.codigo_instalacion or "—"

    # ── Servicios ────────────────────────────────────────────────────────────
    servicios = []
    for sv in (getattr(facility, "servicios", None) or []):
        if isinstance(sv, dict):
            servicios.append({
                "detalle":     sv.get("detalle", "—") or "—",
                "descripcion": sv.get("descripcion", "—") or "—",
            })
    if not servicios:
        servicios = [{"detalle": "—", "descripcion": "—"}]

    # ── Equipos entregados al cliente ─────────────────────────────────────────
    equipos = []
    for it in (facility.items_planificados or []):
        if it.get("destino_final", "cliente") != "cliente":
            continue
        item_obj = None
        try:
            from config.apps.inventory.models.item import Item
            item_obj = Item.objects(id=it.get("item_id")).first()
        except Exception:
            pass
        equipos.append({
            "nombre": (item_obj.nombre if item_obj else None) or "—",
            "serial": str(getattr(item_obj, "serial", None) or "—") if item_obj else "—",
            "estado": ((getattr(item_obj, "estado", None) or "—").replace("_", " ")) if item_obj else "—",
        })
    if not equipos:
        equipos = [{"nombre": "Sin equipos entregados al cliente.", "serial": "—", "estado": "—"}]

    # ── Abrir template y rellenar ─────────────────────────────────────────────
    doc = Document(_TEMPLATE)

    # 1. Carátula: reemplaza subtítulo y borra título original
    _replace_in_section(doc, "JEFATURA NACIONAL TÉCNICA",
                        "ACTA DE ENTREGA DE PRESTACIÓN DE SERVICIOS CORPORATIVOS")
    _replace_in_section(doc, "ACTA DE ENTREGA DE SERVICIO", "")

    # 2. Reemplazar TODAS las fechas estáticas del documento
    _replace_in_section(doc, "27 de abril de 2026", fecha_fin_larga)

    # 2. Tabla 2 — Información General
    #    [0,1]=Cliente  [0,3]=OT  [1,1]=Técnico  [2,1]=Dirección  [3,1]=Fecha
    t2 = doc.tables[2]
    _set_cell(t2.cell(0, 1), cliente_nombre)
    _set_cell(t2.cell(0, 3), ot_code)
    _set_cell(t2.cell(1, 1), tecnico_nombre)
    _set_cell(t2.cell(2, 1), direccion)
    _set_cell(t2.cell(3, 1), fecha_fin_corta)

    # 3. Tabla 3 — Servicios Entregados (filas variables)
    _fill_table_rows(doc.tables[3], servicios, ["detalle", "descripcion"])

    # 4. Tabla 4 — Equipos Entregados (filas variables)
    _fill_table_rows(doc.tables[4], equipos, ["nombre", "serial", "estado"])

    # 5. Párrafo de conformidad
    #    "... a (nombre cliente), el día (fecha fin de instalación)."
    for p in doc.paragraphs:
        if "(nombre cliente)" in p.text and "conformidad" in p.text:
            _replace_in_para(p, "(nombre cliente)", cliente_nombre)
            _replace_in_para(p, "(fecha fin de instalación)", fecha_fin_larga)
            break

    # 6. Tabla 5 — Firmas
    # Solo reemplazamos el texto en los párrafos existentes del template
    # para conservar EXACTAMENTE el mismo formato (espaciado, fuente) que la col 1.
    # Estructura original cell[2,1]: párrafo 0 = nombre, 1 = "Técnico", 2 = "DATACOM"
    # Estructura original cell[2,2]: párrafo 0 = "Representante del Cliente", 1 = nombre cliente
    t5 = doc.tables[5]
    _replace_in_para(t5.cell(2, 1).paragraphs[0],
                     "(nombre del técnico asignado)", tecnico_nombre)
    _replace_in_para(t5.cell(2, 2).paragraphs[1],
                     "(nombre cliente)", cliente_nombre)

    # 7. Mantener la sección de firmas entera en una sola página
    _keep_section_together(doc, "FIRMAS Y CONFORMIDAD", t5)

    # ── Guardar .docx y convertir a PDF ──────────────────────────────────────
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "acta.docx")
        doc.save(docx_path)

        result = subprocess.run(
            [
                "libreoffice", "--headless", "--norestore",
                "--convert-to", "pdf",
                "--outdir", tmpdir,
                docx_path,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice falló (código {result.returncode}):\n{result.stderr}"
            )

        pdf_files = glob.glob(os.path.join(tmpdir, "*.pdf"))
        if not pdf_files:
            raise RuntimeError("LibreOffice no generó el PDF")

        with open(pdf_files[0], "rb") as f:
            return f.read()
